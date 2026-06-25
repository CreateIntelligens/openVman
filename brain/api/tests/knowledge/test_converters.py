"""Tests for knowledge.converters file → plain-text extraction."""

from __future__ import annotations

import importlib
import sys
from pathlib import Path

API_ROOT = Path(__file__).resolve().parents[1]
if str(API_ROOT) not in sys.path:
    sys.path.insert(0, str(API_ROOT))

converters = importlib.import_module("knowledge.converters")


def test_csv_emits_col_value_lines(tmp_path: Path):
    csv_path = tmp_path / "hours.csv"
    csv_path.write_text(
        "星期,上午,下午\n週一,09:00-12:00,14:00-18:00\n週二,休診,14:00-18:00\n",
        encoding="utf-8",
    )

    result = converters.convert_to_text(csv_path)

    assert result is not None
    lines = result.splitlines()
    assert lines[0] == "星期: 週一 ｜ 上午: 09:00-12:00 ｜ 下午: 14:00-18:00"
    assert lines[1] == "星期: 週二 ｜ 上午: 休診 ｜ 下午: 14:00-18:00"


def test_txt_passthrough(tmp_path: Path):
    txt_path = tmp_path / "note.txt"
    content = "第一行\n第二行\n"
    txt_path.write_text(content, encoding="utf-8")

    assert converters.convert_to_text(txt_path) == content


def test_md_passthrough(tmp_path: Path):
    md_path = tmp_path / "doc.md"
    content = "# 標題\n\n內容段落。\n"
    md_path.write_text(content, encoding="utf-8")

    assert converters.convert_to_text(md_path) == content


def test_xlsx_round_trip(tmp_path: Path):
    import openpyxl

    xlsx_path = tmp_path / "book.xlsx"
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "門診"
    sheet.append(["科別", "醫師"])
    sheet.append(["內科", "王醫師"])
    sheet.append(["外科", "李醫師"])
    workbook.save(xlsx_path)

    result = converters.convert_to_text(xlsx_path)

    assert result is not None
    assert "# 門診" in result
    assert "科別: 內科 ｜ 醫師: 王醫師" in result
    assert "科別: 外科 ｜ 醫師: 李醫師" in result


def test_docx_round_trip(tmp_path: Path):
    import docx

    docx_path = tmp_path / "doc.docx"
    document = docx.Document()
    document.add_paragraph("這是第一段。")
    document.add_paragraph("這是第二段。")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "姓名"
    table.cell(0, 1).text = "電話"
    table.cell(1, 0).text = "張三"
    table.cell(1, 1).text = "0912345678"
    document.save(docx_path)

    result = converters.convert_to_text(docx_path)

    assert result is not None
    assert "這是第一段。" in result
    assert "這是第二段。" in result
    assert "姓名" in result
    assert "張三" in result
    assert "0912345678" in result


def test_unsupported_suffix_returns_none(tmp_path: Path):
    zip_path = tmp_path / "archive.zip"
    zip_path.write_bytes(b"PK\x03\x04not-a-real-zip")

    assert converters.convert_to_text(zip_path) is None


def test_corrupt_docx_returns_none(tmp_path: Path):
    bad_path = tmp_path / "broken.docx"
    bad_path.write_bytes(b"this is not a valid docx file at all")

    assert converters.convert_to_text(bad_path) is None


def test_empty_text_returns_none(tmp_path: Path):
    empty_path = tmp_path / "empty.txt"
    empty_path.write_text("   \n\t\n", encoding="utf-8")

    assert converters.convert_to_text(empty_path) is None


def test_supported_suffixes_set():
    assert converters.SUPPORTED_SUFFIXES == frozenset(
        {".md", ".txt", ".csv", ".docx", ".xlsx", ".pdf"}
    )
