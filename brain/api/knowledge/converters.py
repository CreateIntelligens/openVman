"""Extract readable plain text from arbitrary uploaded files.

This module performs *format → text* extraction only. It does NOT produce
final markdown — a separate LLM normalizer cleans the raw text afterwards.

Each format has a small private helper (``_convert_*``). The public entry
point :func:`convert_to_text` dispatches on suffix and never raises: any
extraction failure (unsupported format, corrupt file, empty result) yields
``None`` so the caller can leave such files untouched.
"""

from __future__ import annotations

import csv
import logging
from collections.abc import Sequence
from pathlib import Path

logger = logging.getLogger("brain.knowledge.converters")

# Joins column:value pairs within a single row (full-width separators read
# better for the predominantly Chinese-language knowledge base).
_PAIR_SEP = " ｜ "
_KV_SEP = ": "


def _row_to_line(header: list[str], row: Sequence[object]) -> str:
    """Render one table row as ``col: val ｜ col: val`` using header names.

    Cell values may be non-strings (xlsx yields int/float/None), so each value
    is coerced via ``str`` before trimming.
    """
    pairs: list[str] = []
    for index, value in enumerate(row):
        cell = "" if value is None else str(value).strip()
        if cell == "":
            continue
        column = header[index] if index < len(header) else f"col{index + 1}"
        pairs.append(f"{column}{_KV_SEP}{cell}")
    return _PAIR_SEP.join(pairs)


def _convert_text(path: Path) -> str:
    """Read a plain-text / markdown file verbatim."""
    return path.read_text(encoding="utf-8-sig")


def _convert_csv(path: Path) -> str:
    """Render a CSV as one ``col: val ｜ ...`` line per data row."""
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        reader = csv.reader(handle)
        rows = list(reader)
    if not rows:
        return ""
    header = [str(cell).strip() for cell in rows[0]]
    lines = [_row_to_line(header, row) for row in rows[1:]]
    return "\n".join(line for line in lines if line)


def _convert_docx(path: Path) -> str:
    """Extract paragraph text and table cell text from a .docx document."""
    import docx

    document = docx.Document(str(path))
    parts: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            line = _PAIR_SEP.join(cell for cell in cells if cell)
            if line:
                parts.append(line)

    return "\n".join(parts)


def _convert_xlsx(path: Path) -> str:
    """Render each worksheet as ``# title`` plus ``col: val ｜ ...`` rows."""
    import openpyxl

    workbook = openpyxl.load_workbook(path, read_only=True, data_only=True)
    sections: list[str] = []
    try:
        for sheet in workbook.worksheets:
            rows = list(sheet.iter_rows(values_only=True))
            if not rows:
                continue
            header = [str(cell).strip() if cell is not None else "" for cell in rows[0]]
            lines = [f"# {sheet.title}"]
            for row in rows[1:]:
                line = _row_to_line(header, list(row))
                if line:
                    lines.append(line)
            if len(lines) > 1:
                sections.append("\n".join(lines))
    finally:
        workbook.close()
    return "\n\n".join(sections)


def _convert_pdf(path: Path) -> str:
    """Extract text from each PDF page via PyMuPDF, joined by blank lines."""
    import fitz

    pages: list[str] = []
    with fitz.open(path) as document:
        for page in document:
            # get_text() with no args returns the page's plain text; the stub
            # types it as a union, so coerce to str before stripping.
            text = str(page.get_text()).strip()
            if text:
                pages.append(text)
    return "\n\n".join(pages)


# Suffix → extraction helper. Drives both dispatch and SUPPORTED_SUFFIXES.
_CONVERTERS = {
    ".md": _convert_text,
    ".txt": _convert_text,
    ".csv": _convert_csv,
    ".docx": _convert_docx,
    ".xlsx": _convert_xlsx,
    ".pdf": _convert_pdf,
}

SUPPORTED_SUFFIXES: frozenset[str] = frozenset(_CONVERTERS)


def convert_to_text(path: Path) -> str | None:
    """Extract readable text from a file.

    Returns ``None`` if the format is unsupported or extraction fails
    (caller will leave such files untouched). Never raises.
    """
    converter = _CONVERTERS.get(path.suffix.lower())
    if converter is None:
        return None

    try:
        text = converter(path)
    except Exception:  # noqa: BLE001 — any failure → leave file untouched
        logger.warning("Failed to extract text from %s", path, exc_info=True)
        return None

    if text is None or not text.strip():
        return None
    return text
